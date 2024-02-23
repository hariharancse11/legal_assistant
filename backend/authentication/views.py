# authentication/views.py
from rest_framework import generics, status
from rest_framework.views import APIView
from rest_framework.response import Response
from django.contrib.auth.models import User
from django.contrib.auth import authenticate, login, logout
from rest_framework.permissions import AllowAny, IsAuthenticated
from .serializers import UserSerializer
from .llm_views import bill_summerizer,letter_generator,chat

# views.py

from rest_framework import generics, status
from rest_framework.response import Response
from rest_framework.authtoken.models import Token
from django.contrib.auth import authenticate
from .serializers import UserSerializer

class UserSignUp(generics.CreateAPIView):
    serializer_class = UserSerializer

    def post(self, request, *args, **kwargs):
        serializer = self.get_serializer(data=request.data)
        if serializer.is_valid():
            user = serializer.save()
            token, _ = Token.objects.get_or_create(user=user)
            return Response({'status':True,'token': token.key}, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

class UserSignIn(generics.CreateAPIView):
    serializer_class = UserSerializer

    def post(self, request, *args, **kwargs):
        username = request.data.get('email')
        password = request.data.get('password')
        user = authenticate(username=username, password=password)
        if user:
            token, _ = Token.objects.get_or_create(user=user)
            return Response({'status':True,'message':'Successfully signed in','token': token.key}, status=status.HTTP_200_OK)
        else:
            return Response({'status':False,'error': 'Invalid credentials'}, status=status.HTTP_401_UNAUTHORIZED)
from rest_framework.authentication import TokenAuthentication


class UserLogout(APIView):
    # authentication_classes = [TokenAuthentication]
    # permission_classes = [IsAuthenticated]

    def get(self, request, *args, **kwargs):
        # # Get the current user's token
        # token = request.auth
        # # Delete the token from the database
        # token.delete()
        return Response({'status':True,'message': 'Logged out successfully'}, status=status.HTTP_200_OK)


import os
import tempfile
from rest_framework import status
from rest_framework.parsers import MultiPartParser
from rest_framework.response import Response
from rest_framework.views import APIView
# from .models import File
# from .serializers import FileSerializer
import easyocr
from docx import Document
import PyPDF2
from PyPDF2 import PdfReader


class SummaryView(APIView):
    parser_classes = (MultiPartParser,)

    def post(self, request, *args, **kwargs):
        file_obj = request.FILES.get('file')

        if file_obj.name.lower().endswith(('.png', '.jpg', '.jpeg', '.gif')):
            # If the uploaded file is an image
            print('Image uploaded')
            reader = easyocr.Reader(['en'])  # Language can be changed as needed
            img_path = os.path.join(tempfile.gettempdir(), file_obj.name)
            with open(img_path, 'wb') as f:
                for chunk in file_obj.chunks():
                    f.write(chunk)
            result = reader.readtext(img_path)
            os.remove(img_path)
            print('Extracted text:', result)
            return Response({'summary':bill_summerizer(result)})
        elif file_obj.name.lower().endswith('.docx'):
            # If the uploaded file is a DOCX document
            print('DOCX document uploaded')
            docx_file = Document(file_obj)
            text = ""

            # Extract text from paragraphs
            for paragraph in docx_file.paragraphs:
                text += paragraph.text + "\n"

            # Extract text from tables
            for table in docx_file.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            text += paragraph.text + "\n"

            # Extract text from headers and footers
            for section in docx_file.sections:
                for header in section.header.paragraphs:
                    text += header.text + "\n"
                for footer in section.footer.paragraphs:
                    text += footer.text + "\n"

            # print(text)
            return Response({'summary':bill_summerizer(text)})        
        elif file_obj.name.lower().endswith('.pdf'):
            # If the uploaded file is a PDF document
            print('PDF document uploaded')
            pdf_reader = PdfReader(file_obj)
            text = ''
            for page_num in range(len(pdf_reader.pages)):
                text += pdf_reader.pages[page_num].extract_text()
            print('Extracted text:', text)

            return Response({'summary':bill_summerizer(text)})
        else:
            return Response({'error': 'Unsupported file format'}, status=status.HTTP_400_BAD_REQUEST)

        # serializer = FileSerializer(data={'file': file_obj})
        # if serializer.is_valid():
        #     serializer.save()
        #     return Response(serializer.data, status=status.HTTP_201_CREATED)
        # else:
        #     return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)



class LetterView(APIView):
    parser_classes = (MultiPartParser,)

    def post(self, request, *args, **kwargs):
        file_obj = request.FILES.get('file')
        extra = request.data.get('extra')
        # print(extra)


        if file_obj.name.lower().endswith(('.png', '.jpg', '.jpeg', '.gif')):
            # If the uploaded file is an image
            print('Image uploaded')
            reader = easyocr.Reader(['en'])  # Language can be changed as needed
            img_path = os.path.join(tempfile.gettempdir(), file_obj.name)
            with open(img_path, 'wb') as f:
                for chunk in file_obj.chunks():
                    f.write(chunk)
            result = reader.readtext(img_path)
            os.remove(img_path)
            print('Extracted text:', result)
            return Response({'summary':letter_generator(result.append(extra))})
        elif file_obj.name.lower().endswith('.docx'):
            # If the uploaded file is a DOCX document
            print('DOCX document uploaded')
            docx_file = Document(file_obj)
            text = ""

            # Extract text from paragraphs
            for paragraph in docx_file.paragraphs:
                text += paragraph.text + "\n"

            # Extract text from tables
            for table in docx_file.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            text += paragraph.text + "\n"

            # Extract text from headers and footers
            for section in docx_file.sections:
                for header in section.header.paragraphs:
                    text += header.text + "\n"
                for footer in section.footer.paragraphs:
                    text += footer.text + "\n"
            return Response({'summary':letter_generator(text+' '+extra)})
        elif file_obj.name.lower().endswith('.pdf'):
            # If the uploaded file is a PDF document
            print('PDF document uploaded')
            pdf_reader = PdfReader(file_obj)
            text = ''
            for page_num in range(len(pdf_reader.pages)):
                text += pdf_reader.pages[page_num].extract_text()
            print('Extracted text:', text)

            return Response({'summary':letter_generator(text+' '+extra)})
        else:
            return Response({'error': 'Unsupported file format'}, status=status.HTTP_400_BAD_REQUEST)

        # serializer = FileSerializer(data={'file': file_obj})
        # if serializer.is_valid():
        #     serializer.save()
        #     return Response(serializer.data, status=status.HTTP_201_CREATED)
        # else:
        #     return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
class ChatView(APIView):
    def post(self, request, *args, **kwargs):
        prompt = request.data.get('prompt')
        # print(prompt)
        return Response({'text':chat(prompt)})
    
from django.contrib.auth import update_session_auth_hash

class ChangePasswordView(APIView):
    # permission_classes = (IsAuthenticated,)

    def post(self, request, *args, **kwargs):
        old_password = request.data.get('old_password')
        new_password = request.data.get('new_password')
        email = request.data.get("email")
        try:
            user = User.objects.get(email=email)
            if user:
                if not user.check_password(old_password):
                    return Response({"status": False})
            user.set_password(new_password)
            user.save()
            return Response({"status":True})

        except:
            return Response({"status":False})